"""
gmp_theme.py — Module de thème GMP pour la génération de documents Word (.docx)

Charte graphique GMP (fond blanc) :
  Primaire  : #F29400 (orange GMP)  — titres H1, en-têtes tableaux
  Secondaire: #0A3D67 (bleu GMP)    — titres H2, filets H2, 1ère colonne
  Accent    : #fbbf24 (amber)       — titres H3, numéros
  Texte     : #171717               — corps
  Fond      : #ffffff               — page
  Fond carte: #FFF8EC               — ligne alternée (orange très clair)
  Bordure   : #e2e8f0

Usage :
    from gmp_theme import GmpDoc

    doc = GmpDoc("Mon document", "Sous-titre", "Mars 2026", "Jean Dupont")
    doc.heading("1. Titre de section")
    doc.heading("1.1 Sous-section", level=2)
    doc.body("Paragraphe de texte.")
    doc.bullet("Point de liste")
    doc.table(["Col A", "Col B"], [["val 1", "val 2"]], col_widths=[4, 12])
    doc.save("mon-document.docx")
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette GMP (fond blanc) ─────────────────────────────────────────────────
ORANGE  = RGBColor(0xF2, 0x94, 0x00)   # primaire — titres H1, en-têtes
BLUE    = RGBColor(0x0A, 0x3D, 0x67)   # secondaire — titres H2, 1ère colonne
AMBER   = RGBColor(0xfb, 0xbf, 0x24)   # accent — H3, numéros
TEXT    = RGBColor(0x17, 0x17, 0x17)   # texte corps
MUTED   = RGBColor(0x64, 0x74, 0x8b)   # texte atténué
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = RGBColor(0xFF, 0xF8, 0xEC)   # fond ligne alternée (orange très clair)
BORDER  = RGBColor(0xe2, 0xe8, 0xf0)   # bordure légère

FONT = "Outfit"

# ── Helpers internes ─────────────────────────────────────────────────────────

def _hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _hex(rgb))
    tcPr.append(shd)


def _add_border_bottom(paragraph, color: str, size: str = "12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _spacing(paragraph, before: int = 0, after: int = 120):
    pPr = paragraph._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"),  str(after))
    pPr.append(sp)


def _style_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name   = FONT
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color or TEXT


def _set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ── Classe principale ─────────────────────────────────────────────────────────

class GmpDoc:
    """
    Générateur de documents Word aux couleurs de la charte GMP.

    Paramètres
    ----------
    title       : Titre principal (page de garde + en-tête)
    subtitle    : Sous-titre (page de garde)
    date        : Date du document (ex: "Mars 2026")
    authors     : Auteur(s) (ex: "Yahaya Coulibaly, Hyosua Colléter")
    version     : Version (défaut "1.0")
    status      : Statut (défaut "Document initial")
    """

    def __init__(
        self,
        title: str,
        subtitle: str = "",
        date: str = "",
        authors: str = "",
        version: str = "1.0",
        status: str = "Document initial",
    ):
        self._doc = Document()
        self._setup_document()
        self._build_cover(title, subtitle, date, authors, version, status)
        self._doc.add_page_break()

    # ── Setup ────────────────────────────────────────────────────────────────

    def _setup_document(self):
        s = self._doc.sections[0]
        s.left_margin = s.right_margin = Cm(2.5)
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.0)
        normal = self._doc.styles["Normal"]
        normal.font.name      = FONT
        normal.font.size      = Pt(11)
        normal.font.color.rgb = TEXT

    # ── Page de garde ────────────────────────────────────────────────────────

    def _build_cover(self, title, subtitle, date, authors, version, status):
        doc = self._doc

        # Bandeau titre
        banner = doc.add_table(rows=1, cols=1)
        banner.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = banner.rows[0].cells[0]
        _set_cell_bg(cell, ORANGE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(p.add_run(title.upper()), size=24, bold=True, color=WHITE)
        _spacing(p, before=220, after=220)

        _spacing(doc.add_paragraph(), after=80)

        if subtitle:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _style_run(p2.add_run(subtitle), size=14, bold=True)
            _spacing(p2, after=60)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(p3.add_run("IUT d'Évry-Courcouronnes — Département GMP"), size=11, color=MUTED)
        _spacing(p3, after=200)

        # Tableau de métadonnées
        meta_data = []
        if version: meta_data.append(("Version", version))
        if date:    meta_data.append(("Date",    date))
        if authors: meta_data.append(("Auteurs", authors))
        if status:  meta_data.append(("Statut",  status))

        if meta_data:
            meta = doc.add_table(rows=len(meta_data), cols=2)
            meta.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (k, v) in enumerate(meta_data):
                c0, c1 = meta.rows[i].cells
                bg = ROW_ALT if i % 2 == 0 else WHITE
                _set_cell_bg(c0, ROW_ALT)
                _set_cell_bg(c1, bg)
                _style_run(c0.paragraphs[0].add_run(k), size=10, bold=True, color=ORANGE)
                _style_run(c1.paragraphs[0].add_run(v), size=10)
            _set_col_width(meta, 0, 3.5)
            _set_col_width(meta, 1, 10.5)

    # ── API publique ─────────────────────────────────────────────────────────

    def heading(self, text: str, level: int = 1) -> None:
        """
        Ajoute un titre.
        level 1 → orange, 22pt, filet bas orange (grandes sections)
        level 2 → bleu GMP, 15pt, filet bas gris (sous-sections)
        level 3 → amber, 12pt, gras (sous-sous-sections)
        """
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        if level == 1:
            _style_run(run, size=22, bold=True, color=ORANGE)
            _add_border_bottom(p, _hex(ORANGE), "16")
            _spacing(p, before=300, after=160)
        elif level == 2:
            _style_run(run, size=15, bold=True, color=BLUE)
            _add_border_bottom(p, _hex(BORDER), "8")
            _spacing(p, before=200, after=100)
        elif level == 3:
            _style_run(run, size=12, bold=True, color=AMBER)
            _spacing(p, before=140, after=60)

    def body(self, text: str) -> None:
        """Ajoute un paragraphe de texte courant."""
        p = self._doc.add_paragraph()
        _style_run(p.add_run(text))
        _spacing(p, after=80)

    def bullet(self, text: str) -> None:
        """Ajoute un point de liste (tiret em)."""
        p = self._doc.add_paragraph()
        _style_run(p.add_run(f"— {text}"))
        p.paragraph_format.left_indent = Cm(0.6)
        _spacing(p, after=60)

    def numbered(self, items: list[tuple[str, str]]) -> None:
        """
        Liste numérotée avec titre et description.
        items = [(titre, description), ...]
        """
        for i, (titre, desc) in enumerate(items, 1):
            p = self._doc.add_paragraph()
            _style_run(p.add_run(f"{i}. "), size=11, bold=True, color=AMBER)
            _style_run(p.add_run(f"{titre} — "), size=11, bold=True)
            _style_run(p.add_run(desc), size=11)
            p.paragraph_format.left_indent = Cm(0.5)
            _spacing(p, after=80)

    def table(
        self,
        headers: list[str],
        rows: list[list[str]],
        col_widths: list[float] | None = None,
    ) -> None:
        """
        Ajoute un tableau stylisé GMP.
        En-tête : fond orange, texte blanc.
        Lignes  : alternance blanc / orange très clair.
        1ère colonne : texte bleu GMP, gras.
        col_widths : largeurs en cm pour chaque colonne.
        """
        t = self._doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.style = "Table Grid"

        # En-tête
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _set_cell_bg(cell, ORANGE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _style_run(cell.paragraphs[0].add_run(h), size=10, bold=True, color=WHITE)

        # Données
        for r_idx, row_data in enumerate(rows):
            bg = ROW_ALT if r_idx % 2 == 0 else WHITE
            for c_idx, val in enumerate(row_data):
                cell = t.rows[r_idx + 1].cells[c_idx]
                _set_cell_bg(cell, bg)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                color = BLUE if c_idx == 0 else TEXT
                _style_run(cell.paragraphs[0].add_run(str(val)), size=10, bold=(c_idx == 0), color=color)

        if col_widths:
            for i, w in enumerate(col_widths):
                _set_col_width(t, i, w)

        _spacing(self._doc.add_paragraph(), after=100)

    def page_break(self) -> None:
        """Insère un saut de page."""
        self._doc.add_page_break()

    def spacer(self, after: int = 100) -> None:
        """Insère un espace vertical."""
        _spacing(self._doc.add_paragraph(), after=after)

    def save(self, path: str) -> None:
        """Sauvegarde le document au chemin indiqué."""
        self._doc.save(path)
        print(f"Généré : {path}")

"""Build charte-graphique.docx — GMP IUT d'Évry"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ───────────────────────────────────────────────────────────────────
DARK_PRIMARY  = RGBColor(0xF2, 0x94, 0x00)   # Orange GMP
LIGHT_PRIMARY = RGBColor(0x0D, 0x94, 0x88)   # Teal

# Swatches (pour les carrés de couleur)
DARK_BG       = RGBColor(0x0A, 0x3D, 0x67)
DARK_CARD     = RGBColor(0x0D, 0x4E, 0x80)
DARK_DEEP     = RGBColor(0x07, 0x2D, 0x4D)
DARK_TEXT     = RGBColor(0xE2, 0xE8, 0xF0)
DARK_BORDER   = RGBColor(0x1A, 0x5A, 0x8A)
DARK_MUTED    = RGBColor(0x64, 0x74, 0x8B)
LIGHT_BG      = RGBColor(0xF8, 0xFA, 0xFC)
LIGHT_MUTED   = RGBColor(0x81, 0xAD, 0xC8)
LIGHT_ACCENT  = RGBColor(0xF5, 0x9E, 0x0B)
LIGHT_BORDER  = RGBColor(0xE2, 0xE8, 0xF0)
AMBER         = RGBColor(0xFB, 0xBF, 0x24)

# Palette document (fond blanc)
BLACK         = RGBColor(0x17, 0x17, 0x17)
GRAY_TEXT     = RGBColor(0x44, 0x55, 0x66)
GRAY_MUTED    = RGBColor(0x9C, 0xA3, 0xAF)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

# Tableaux clairs
TH_BG         = RGBColor(0x0D, 0x94, 0x88)   # Teal header
TH_TEXT       = WHITE
ROW_ALT       = RGBColor(0xF0, 0xFD, 0xFA)   # teal très pâle
ROW_PLAIN     = WHITE
BORDER_COLOR  = RGBColor(0xCC, 0xE8, 0xE4)

# Section title bg
SECTION_BG    = RGBColor(0xF0, 0xFD, 0xFA)

# Fonts
F_BODY  = "Outfit"
F_MONO  = "Geist Mono"

# ── Helpers ───────────────────────────────────────────────────────────────────
def rgb_hex(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

def set_cell_bg(cell, color: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(color))
    tcPr.append(shd)

def set_cell_border(cell, color: RGBColor, sides=("top","bottom","left","right")):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), rgb_hex(color))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def cell_run(cell, text, bold=False, italic=False, color=None, size=9, font=F_BODY):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color

def add_para_bg(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(color))
    pPr.append(shd)

def add_left_bar(para, color: RGBColor, sz=36):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), rgb_hex(color))
    pBdr.append(left)
    pPr.append(pBdr)

def spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)

def indent(para, left=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    pPr.append(ind)

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = F_BODY
style.font.size = Pt(10)
style.font.color.rgb = BLACK

# ── Helpers de contenu ─────────────────────────────────────────────────────────
def section_title(text):
    p = doc.add_paragraph()
    add_para_bg(p, SECTION_BG)
    add_left_bar(p, LIGHT_PRIMARY, sz=36)
    r = p.add_run(f"  {text}")
    r.font.name = F_BODY
    r.font.size = Pt(13)
    r.font.color.rgb = LIGHT_PRIMARY
    r.bold = True
    spacing(p, before=220, after=80)

def kv(key, value, key_color=LIGHT_PRIMARY):
    p = doc.add_paragraph()
    spacing(p, before=30, after=30)
    indent(p, 180)
    rk = p.add_run(f"{key}  ")
    rk.font.name = F_BODY
    rk.font.size = Pt(10)
    rk.font.color.rgb = key_color
    rk.bold = True
    rv = p.add_run(value)
    rv.font.name = F_BODY
    rv.font.size = Pt(10)
    rv.font.color.rgb = BLACK

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(len(rows)+1, len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        set_cell_bg(c, TH_BG)
        set_cell_border(c, TH_BG)
        cell_run(c, h, bold=True, color=TH_TEXT, size=9)
    # Rows
    for ri, row in enumerate(rows):
        bg = ROW_ALT if ri % 2 == 0 else ROW_PLAIN
        for ci, cell_data in enumerate(row):
            c = t.cell(ri+1, ci)
            set_cell_bg(c, bg)
            set_cell_border(c, BORDER_COLOR)
            if isinstance(cell_data, tuple):
                # (text, color, bold, italic, font)
                text, color, bold, italic, font = cell_data
                cell_run(c, text, bold=bold, italic=italic, color=color, size=9, font=font)
            else:
                cell_run(c, str(cell_data), color=BLACK, size=9)
    return t

# ── COVER ─────────────────────────────────────────────────────────────────────
cover = doc.add_table(1, 1)
cover.style = "Table Grid"
cell = cover.cell(0, 0)
set_cell_bg(cell, LIGHT_PRIMARY)
for side in ("top","bottom","left","right"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "none")
    tcBorders.append(el)
    tcPr.append(tcBorders)

cell.text = ""
p = cell.add_paragraph()
r = p.add_run("IUT D'ÉVRY — DÉPARTEMENT GMP")
r.font.name = F_BODY; r.font.size = Pt(9); r.font.color.rgb = WHITE; r.bold = True
spacing(p, before=140, after=40)

p = cell.add_paragraph()
r = p.add_run("CHARTE GRAPHIQUE")
r.font.name = F_BODY; r.font.size = Pt(28); r.font.color.rgb = WHITE; r.bold = True
spacing(p, after=60)

p = cell.add_paragraph()
r = p.add_run("Site institutionnel · Design System")
r.font.name = F_BODY; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC,0xF0,0xEB); r.italic = True
spacing(p, after=140)

doc.add_paragraph()

# ── 01 · IDENTITÉ ─────────────────────────────────────────────────────────────
section_title("01 · IDENTITÉ")
kv("Public cible", "Lycéens (recrutement) · Étudiants (espace connecté) · Entreprises")
kv("Ton", "Institutionnel moderne — crédible sans être poussiéreux")
kv("Dark mode", "Activé par défaut · Persisté via cookie HTTP (theme=dark|light)")

doc.add_paragraph()

# ── 02 · THÈMES ───────────────────────────────────────────────────────────────
section_title("02 · THÈMES")

make_table(
    ["Thème", "Fond principal", "Primaire", "Description"],
    [
        ["Light — V2 Forge Outfit", "#f8fafc",
         ("█ #0d9488  Teal", LIGHT_PRIMARY, True, False, F_BODY),
         "Institutionnel clair"],
        ["Dark — Forge", "#0A3D67",
         ("█ #F29400  Orange GMP", DARK_PRIMARY, True, False, F_BODY),
         "Industriel sombre"],
    ]
)
doc.add_paragraph()

# ── 03 · COULEURS DARK ────────────────────────────────────────────────────────
section_title("03 · COULEURS — DARK (Forge)")

def color_rows(data):
    rows = []
    for role, var, hexval, swatch in data:
        rows.append([
            role,
            (var, LIGHT_PRIMARY, False, True, F_MONO),
            hexval,
            ("██████", swatch, False, False, F_BODY),
        ])
    return rows

dark_colors = [
    ("Fond",         "--c-bg",        "#0A3D67",  DARK_BG),
    ("Fond profond", "--c-bg-deep",   "#072d4d",  DARK_DEEP),
    ("Fond carte",   "--c-bg-card",   "#0D4E80",  DARK_CARD),
    ("Primaire",     "--c-primary",   "#F29400",  DARK_PRIMARY),
    ("Texte",        "--c-secondary", "#e2e8f0",  DARK_TEXT),
    ("Accent",       "--c-accent",    "#fbbf24",  AMBER),
    ("Bordure",      "--c-border",    "#1A5A8A",  DARK_BORDER),
    ("Neutre",       "--c-muted",     "#64748b",  DARK_MUTED),
]
make_table(["Rôle", "Variable CSS", "Valeur", "Aperçu"], color_rows(dark_colors))
doc.add_paragraph()

# ── 04 · COULEURS LIGHT ───────────────────────────────────────────────────────
section_title("04 · COULEURS — LIGHT (V2 Forge Outfit)")

light_colors = [
    ("Fond",         "--c-bg",        "#f8fafc",  LIGHT_BG),
    ("Fond profond", "--c-bg-deep",   "#f1f5f9",  RGBColor(0xF1,0xF5,0xF9)),
    ("Fond carte",   "--c-bg-card",   "#ffffff",  WHITE),
    ("Primaire",     "--c-primary",   "#0d9488",  LIGHT_PRIMARY),
    ("Texte",        "--c-secondary", "#171717",  BLACK),
    ("Accent",       "--c-accent",    "#f59e0b",  LIGHT_ACCENT),
    ("Bordure",      "--c-border",    "#e2e8f0",  LIGHT_BORDER),
    ("Neutre",       "--c-muted",     "#81adc8",  LIGHT_MUTED),
]
make_table(["Rôle", "Variable CSS", "Valeur", "Aperçu"], color_rows(light_colors))
doc.add_paragraph()

# ── 05 · VARIABLES UTILITAIRES ────────────────────────────────────────────────
section_title("05 · VARIABLES UTILITAIRES (opacité)")

make_table(
    ["Variable", "Usage"],
    [
        [("--c-primary-15 / -20 / -30", LIGHT_PRIMARY, False, True, F_MONO), "Fonds primaire translucide"],
        [("--c-accent-30 / -40", LIGHT_PRIMARY, False, True, F_MONO), "Fonds accent translucide"],
        [("--c-secondary-25 / -30 / -40 / -80 / -90", LIGHT_PRIMARY, False, True, F_MONO), "Texte secondaire atténué"],
        [("--c-photo-overlay", LIGHT_PRIMARY, False, True, F_MONO), "Overlay couleur sur photos N&B"],
        [("--c-grid-color", LIGHT_PRIMARY, False, True, F_MONO), "Grille blueprint"],
        [("--c-primary-hover", LIGHT_PRIMARY, False, True, F_MONO), "Hover bouton primaire"],
    ]
)

p_rule = doc.add_paragraph()
add_para_bg(p_rule, RGBColor(0xFF, 0xF7, 0xE6))
add_left_bar(p_rule, DARK_PRIMARY, sz=24)
spacing(p_rule, before=100, after=80)
indent(p_rule, 120)
r = p_rule.add_run("  Règle : ")
r.font.name = F_BODY; r.font.color.rgb = DARK_PRIMARY; r.bold = True; r.font.size = Pt(10)
r2 = p_rule.add_run("Ne jamais concaténer un hex alpha à une var() — utiliser ces variables.")
r2.font.name = F_BODY; r2.font.color.rgb = BLACK; r2.font.size = Pt(10)

doc.add_paragraph()

# ── 06 · TYPOGRAPHIE ──────────────────────────────────────────────────────────
section_title("06 · TYPOGRAPHIE")

make_table(
    ["Usage", "Police", "Source"],
    [
        ["Corps & titres",              ("Outfit", LIGHT_PRIMARY, True, False, F_BODY),   "Google Fonts"],
        ["Données, codes, labels mono", ("Geist Mono", LIGHT_PRIMARY, True, False, F_MONO), "Next.js built-in"],
    ]
)
kv("Scale", "clamp() pour les titres principaux · Tailwind pour le reste")
kv("Règle", "Pas de 3ème police")

doc.add_paragraph()

# ── 07 · FORMES & ESPACEMENT ──────────────────────────────────────────────────
section_title("07 · FORMES & ESPACEMENT")

make_table(
    ["Élément", "Valeur"],
    [
        ["Border radius cards/boutons", ("0 — angles droits", BLACK, True, False, F_BODY)],
        ["Border radius badges",        ("rounded-full", BLACK, True, False, F_BODY)],
        ["Profondeur",                  "Bordures 1px · pas d'ombres sauf hover"],
        ["Hover boutons",               "box-shadow: 3-4px offset colorée"],
        ["Layout container",            ("max-w-[1280px] mx-auto px-4 md:px-8", GRAY_TEXT, False, True, F_MONO)],
    ]
)
doc.add_paragraph()

# ── 08 · RESPONSIVE ───────────────────────────────────────────────────────────
section_title("08 · RESPONSIVE")
kv("Approche",    "Mobile first")
kv("Breakpoints", "md (768px) · lg (1024px)")
kv("Grilles",     "1 col mobile → 2 col tablette → 3 col desktop")
kv("Hero photo",  "Masquée sur mobile (hidden lg:block)")

doc.add_paragraph()

# ── 09 · NAVIGATION ───────────────────────────────────────────────────────────
section_title("09 · NAVIGATION")
kv("Topbar publique",  "Accueil · Formation · Entreprises · Actualités + ThemeToggle + Connexion")
kv("Mobile",           "Liens masqués (hidden md:flex)")
kv("Barre secondaire", "Apparaît une fois connecté · contextuelle au rôle")

doc.add_paragraph()

# ── 10 · ICONOGRAPHIE & VISUELS ───────────────────────────────────────────────
section_title("10 · ICONOGRAPHIE & VISUELS")
kv("Icônes",       "Lucide React — tree-shakeable, flat/sharp")
kv("Illustrations","Blueprint SVG (presse hydraulique, CNC, engrenages, chaîne) — couleurs via CSS vars")
kv("Photos",       "Ateliers/TP avec overlay monochrome couleur thème (--c-photo-overlay)")
kv("Hero",         "Double cadre offset (bordure primaire + ombre décalée) · coins L-marks · badge référence")

doc.add_paragraph()

# ── 11 · PRINCIPES DIRECTEURS ─────────────────────────────────────────────────
section_title("11 · PRINCIPES DIRECTEURS")

principes = [
    ("Sharp over rounded",        "Les angles droits évoquent la mécanique de précision"),
    ("Borders over shadows",      "Esthétique blueprint technique, flat et net"),
    ("Teal / Orange dominant",    "Teal en light · Orange GMP (#F29400) en dark"),
    ("Amber rare",                "Réservé aux CTA et alertes importantes"),
    ("Bleu acier comme respiration", "#81adc8 pour les zones secondaires"),
    ("Pas de hex alpha",          "Utiliser les variables --c-* — jamais ${C.primary}20"),
]
for i, (titre, detail) in enumerate(principes):
    p = doc.add_paragraph()
    add_para_bg(p, ROW_ALT if i % 2 == 0 else ROW_PLAIN)
    spacing(p, before=50, after=50)
    indent(p, 180)
    r1 = p.add_run(f"{i+1:02d}.  ")
    r1.font.name = F_BODY; r1.font.color.rgb = LIGHT_PRIMARY; r1.font.size = Pt(10); r1.bold = True
    r2 = p.add_run(f"{titre}  ")
    r2.font.name = F_BODY; r2.font.color.rgb = BLACK; r2.font.size = Pt(10); r2.bold = True
    r3 = p.add_run(f"— {detail}")
    r3.font.name = F_BODY; r3.font.color.rgb = GRAY_TEXT; r3.font.size = Pt(10); r3.italic = True

doc.add_paragraph()

# ── 12 · USAGE EN CODE ────────────────────────────────────────────────────────
section_title("12 · USAGE EN CODE")

CODE_BG      = RGBColor(0xF4, 0xF6, 0xF8)
CODE_COMMENT = RGBColor(0x6B, 0x72, 0x80)
CODE_KW      = LIGHT_PRIMARY
CODE_TEXT    = RGBColor(0x1F, 0x2D, 0x3D)

for line in [
    'import { C } from "@/lib/forge";',
    '',
    '// Fond et texte',
    'style={{ background: C.bg, color: C.secondary }}',
    '',
    '// Bordure primaire',
    'style={{ border: `1px solid ${C.primary}` }}',
    '',
    '// Opacité (variable CSS, pas de hex alpha)',
    'style={{ background: "var(--c-primary-15)" }}',
]:
    p = doc.add_paragraph()
    add_para_bg(p, CODE_BG)
    spacing(p, before=18, after=18)
    indent(p, 240)
    if not line:
        continue
    r = p.add_run(line)
    r.font.name = F_MONO
    r.font.size = Pt(8.5)
    if line.startswith("//"):
        r.font.color.rgb = CODE_COMMENT; r.italic = True
    elif line.startswith("import"):
        r.font.color.rgb = CODE_KW
    else:
        r.font.color.rgb = CODE_TEXT

doc.add_paragraph()

# ── Footer ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
spacing(p, before=200)
r = p.add_run("GMP IUT d'Évry  ·  Design System v1  ·  2025")
r.font.name = F_BODY; r.font.size = Pt(8); r.font.color.rgb = GRAY_MUTED
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ──────────────────────────────────────────────────────────────────────────────
out = "/home/hyo/hyosua/Projects/gmp/docs/charte-graphique.docx"
doc.save(out)
print(f"Saved → {out}")
